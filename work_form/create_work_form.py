from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import datetime
import calendar
import pandas as pd
import sys

day_to_chinese = {1: '一', 2: '二', 3: '三', 4: '四', 5: '五'}
MINIMUM_HOUR = 4


class WeekFreeTime:
    def __init__(self, curriculum):
        self.free_time_list = []
        self.curriculum = curriculum
        for date in range(1, 6):
            column = TableGenerator.convert_day_to_chinese(date)
            free_time = WeekFreeTime.calculate_day_free_time(self.curriculum[column])
            self.free_time_list.append(free_time)

    def get_free_time(self, date):
        return self.free_time_list[date.weekday()]

    @staticmethod
    def calculate_day_free_time(curriculum_col):
        free_time = []
        start = 0
        i = 0
        while i != len(curriculum_col):
            if curriculum_col[i] == 1:
                end = i
                free_time.append([start, end])
                while i != len(curriculum_col) and curriculum_col[i] == 1:
                    i += 1
                start = i
            else:
                i += 1
        return free_time


class TableGenerator:
    def __init__(self, curriculum, begin_date, end_date):
        self.curriculum = curriculum
        self.beg_date = begin_date
        self.end_date = end_date

    def generate_table(self, hour, docx, month):
        table = docx.tables[0]
        nowrow = 1
        week_free_times = WeekFreeTime(self.curriculum)
        for date in self.iter_day():
            if date.weekday() < 5:  # 是否為平日
                day_free_times = week_free_times.get_free_time(date)
                for start, end in day_free_times:
                    hour_count = end - start
                    if hour < hour_count or hour_count >= MINIMUM_HOUR:
                        if hour < hour_count:
                            end = start + hour
                        hour_record_count = \
                            TableGenerator.add_valid_hour(table, nowrow, date, start, end)
                        hour -= hour_record_count
                        if hour_record_count != 0:
                            nowrow += 1
        return hour == 0

    def iter_day(self):
        while self.beg_date < self.end_date:
            yield self.beg_date
            self.beg_date += datetime.timedelta(1)


    @staticmethod
    def convert_hour_count_to_valid_hour_count(hour_count):
        valid_hour_count_list = [0, 1, 2, 4, 8]
        for i in range(len(valid_hour_count_list)-1, -1, -1):
            if hour_count >= valid_hour_count_list[i]:
                return valid_hour_count_list[i]

    @staticmethod
    def add_valid_hour(table, nowrow, date, start, end):
        hour_count = end - start
        valid_hour_count = TableGenerator.convert_hour_count_to_valid_hour_count(hour_count)
        end = start + valid_hour_count
        if start != end:
            return TableGenerator.add_hour_record(table, nowrow, date, start, end)
        else:
            return 0

    @staticmethod
    def add_hour_record(table, nowrow, date, start, end):
        hour_count = end - start
        if hour_count > 4:
            hour_count -= 1  # 休息的那一小時
        text_list = [
                     f'{date.year - 1911}/{date.month}/{date.day}',
                     TableGenerator.convert_day_to_chinese(date.isoweekday()),
                     TableGenerator.convert_free_time_interval_to_output_format(start, end),
                     str(hour_count),
                     '資料整理',
                    ]
        for i in range(5):
            TableGenerator.table_run_add_and_set(table, nowrow, i, text_list[i])
        return hour_count

    @staticmethod
    def convert_free_time_interval_to_output_format(start, end):
        s = TableGenerator.convert_index_to_curse_start_time(start)
        e = TableGenerator.convert_index_to_curse_start_time(end)
        hour_count = end-start
        output = s + '~' + e
        if hour_count > 4:
            rest_time = TableGenerator.convert_index_to_curse_start_time(start+4) + '-' + \
                        TableGenerator.convert_index_to_curse_start_time(start+5)
            output += '\n(' + rest_time + '休息)'
        return output

    @staticmethod
    def convert_index_to_curse_start_time(idx):
        return f'{idx+8}:00'

    @staticmethod
    def table_run_add_and_set(table, row, col, text):
        run = table.cell(row, col).paragraphs[0].add_run(text)
        table.cell(row, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平置中
        TableGenerator.process_chinese_setting(run)

    @staticmethod
    def text_run_add_and_set(docx, idx, pre, hour, suf):
        docx.paragraphs[idx].text = ''
        run = docx.paragraphs[idx].add_run(pre)
        TableGenerator.process_run_setting(run)
        run = docx.paragraphs[idx].add_run(str(hour))
        TableGenerator.process_run_setting(run)
        run.underline = True
        run = docx.paragraphs[idx].add_run(suf)
        TableGenerator.process_run_setting(run)

    @staticmethod
    def process_run_setting(run):
        TableGenerator.process_chinese_setting(run)
        run.font.size = Pt(14)
        run.bold = True

    @staticmethod
    def process_chinese_setting(run):
        run.font.name = u'標楷體'
        r = run._element.rPr.rFonts  # 中文特有的處理
        r.set(qn("w:eastAsia"), "標楷體")

    @staticmethod
    def convert_day_to_chinese(day):
        return day_to_chinese.get(day, 'None')


class DateAdapter:
    def __init__(self, date, default_value):
        self.date = date
        self.value = default_value

    def get_date(self):
        try:
            date = datetime.date(datetime.date.today().year, month, int(self.date))
        except ValueError:
            if self.date == "":
                date = datetime.date(datetime.date.today().year, month, self.value)
            else:
                print('invalid date!')
                sys.exit(1)
        return date


docx = Document("work.docx")
df = pd.read_excel("work.xlsx", engine='openpyxl')
df = df.append({"一": 1, "二": 1, "三": 1, "四": 1, "五": 1}, ignore_index=True)

hour = int(input("請輸入工時: "))
month = int(input("請輸入月份: "))

beg = input("請輸入工讀開始日(如不輸入則預設為月份的第一天): ")
end = input("請輸入工讀結束日(如不輸入則預設為月份的最後一天): ")
beg_date = DateAdapter(beg, 1).get_date()
end_date = DateAdapter(end, calendar.mdays[month]).get_date()

table_generator = TableGenerator(df, beg_date, end_date)

if table_generator.generate_table(hour, docx, month):
    print('產生word完成。')

    TableGenerator.text_run_add_and_set(docx, 2, '*正常工作時數：', f'      {hour}     ', '小時')
    TableGenerator.text_run_add_and_set(docx, 3, '*薪資小計：', f'      {hour*160}      ', '元(A)')
    TableGenerator.text_run_add_and_set(docx, 5, '*合計應領薪資：', f'      {hour*160}     ', '元(A+B) /　工讀生指導人：______________')

    docx.save(str(month) + '月份工作表.docx')
else:
    print('工作日不夠，產生word失敗。')
